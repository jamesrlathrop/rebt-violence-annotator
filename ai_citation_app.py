import streamlit as st
import pdfplumber
import docx
from docx import Document
import io
import os

# Dummy citation lookup function (replace with real lookup)
def find_apa_citation(query):
    return f"(Author, Year)"  # Replace this with your citation logic

def insert_citations(text):
    import re
    def replace_fn(match):
        query = match.group(1).strip()
        citation = find_apa_citation(query)
        return f"{query} {citation}"
    return re.sub(r"\[\[CITE:(.*?)\]\]", replace_fn, text)

def convert_to_docx(text, filename):
    doc = Document()
    for para in text.split('\n'):
        doc.add_paragraph(para)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

st.title("AI Research Citation Assistant")

uploaded_file = st.file_uploader("Upload a .docx manuscript", type="docx")

if uploaded_file:
    doc = docx.Document(uploaded_file)
    full_text = "\n".join(para.text for para in doc.paragraphs)

    st.subheader("Original Text")
    st.text_area("View your manuscript", value=full_text, height=300)

    if st.button("Generate APA Citations"):
        cited_text = insert_citations(full_text)

        st.subheader("Text with Inline Citations")
        st.text_area("Output", value=cited_text, height=300)

        # Word download
        docx_bytes = convert_to_docx(cited_text, "output_with_citations.docx")
        st.download_button(
            label="Download .docx with Citations",
            data=docx_bytes,
            file_name="output_with_citations.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
